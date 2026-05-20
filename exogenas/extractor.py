"""Extractor de certificados de retención en la fuente (Formato 1003)."""

from __future__ import annotations

import re
from pathlib import Path

import pdfplumber

from exogenas.municipios import buscar_municipio

# ── Helpers numéricos ─────────────────────────────────────────────────────────

def _parse_money(s: str) -> float:
    """Convierte cadena monetaria colombiana o US a float.
    Soporta: '1.234.567,00', '1,234,567.00', '$1.234', '($ 144.325)', '3.622.900'.
    """
    s = re.sub(r"[\$\(\)\s]", "", str(s)).strip()
    if not s:
        return 0.0
    # Rechazar fechas: "25.12.31", "2025.12.31", "31/12/2025", "31-12-25"
    if re.match(r"^\d{2,4}[.\-/]\d{1,2}[.\-/]\d{2,4}$", s):
        return 0.0
    dot_count   = s.count(".")
    comma_count = s.count(",")

    if dot_count > 1:
        # Múltiples puntos → todos son miles: 3.622.900 o 1.234.567,00
        s = s.replace(".", "")
        if comma_count == 1:
            s = s.replace(",", ".")
    elif comma_count > 1:
        # Múltiples comas → miles: 1,234,567
        s = s.replace(",", "")
    elif dot_count == 1 and comma_count == 1:
        if s.rfind(",") > s.rfind("."):   # 1.234,56 → decimal coma
            s = s.replace(".", "").replace(",", ".")
        else:                              # 1,234.56 → decimal punto
            s = s.replace(",", "")
    elif dot_count == 1:
        # Un solo punto: si tiene exactamente 3 decimales → miles (144.325 → 144325)
        # si tiene 1-2 decimales → decimal real (2.50 → 2.50)
        parts = s.split(".")
        if len(parts[1]) == 3:
            s = s.replace(".", "")        # 144.325 → 144325
        # else: float as-is (2.50, 0.70 etc.)
    elif comma_count == 1:
        parts = s.split(",")
        if len(parts[1]) <= 2:
            s = s.replace(",", ".")       # 144,32 → 144.32
        else:
            s = s.replace(",", "")        # 1,234 → 1234
    try:
        return float(s)
    except ValueError:
        return 0.0


def calcular_dv(nit: str) -> str:
    """Dígito de verificación colombiano (algoritmo DIAN)."""
    weights = [71, 67, 59, 53, 47, 43, 41, 37, 29, 23, 19, 17, 13, 7, 3]
    digits = re.sub(r"\D", "", nit)
    if not digits:
        return ""
    n = len(digits)
    w = weights[-n:] if n <= len(weights) else weights
    total = sum(int(d) * wt for d, wt in zip(digits, w))
    r = total % 11
    return str(r) if r in (0, 1) else str(11 - r)


def _tipo_doc(nit: str) -> str:
    return "31" if len(re.sub(r"\D", "", nit)) >= 9 else "13"


def _tipo_doc_from_context(nit: str, text: str) -> str:
    """Determina tipo_doc buscando el label explícito antes del NIT en el texto."""
    digits = re.sub(r"\D", "", nit)
    if not digits:
        return "31"
    # Buscar la posición del NIT en el texto
    pos = text.find(digits[:9])
    if pos > 0:
        context = text[max(0, pos - 150): pos + 20]
        if re.search(r"\bC\.?C\.?\b|\bc[eé]dula\b|\bidentificaci[oó]n\b", context, re.I):
            if not re.search(r"\bN\.?I\.?T\.?\b", context[-60:], re.I):
                return "13"
        if re.search(r"\bN\.?I\.?T\.?\b", context, re.I):
            return "31"
    return _tipo_doc(nit)


# ── Regex patterns ────────────────────────────────────────────────────────────

# NIT con/sin DV — separador obligatorio ante DV para evitar backtracking
_RE_NIT_DV   = re.compile(
    r"N\.?I\.?T\.?\s*[:\s#No.]*([0-9]{6,12}|[0-9]{1,3}(?:\.[0-9]{3})+)\s*[-–]\s*([0-9])\b",
    re.I,
)
_RE_NIT_NODV = re.compile(
    r"N\.?I\.?T\.?\s*[:\s#No.]*([0-9]{6,12}|[0-9]{1,3}(?:\.[0-9]{3})+)",
    re.I,
)

# Tipo de certificado
_RE_CERT_ICA   = re.compile(
    r"retenci[oó]n\s+(?:de\s+|por\s+|y\s+)?ica\b"
    r"|ica\s+\d+[xX]\d+"          # "ICA 7x1000"
    r"|certificado\s+de\s+retenci[oó]n\s+(?:por\s+)?ica\b",
    re.I,
)
_RE_CERT_IVA   = re.compile(r"retenci[oó]n(?:es)?\s+(?:de\s+|por\s+)?IVA|RETEIVA", re.I)
_RE_CERT_RENTA = re.compile(
    r"retenci[oó]n(?:es)?\s+(?:en\s+la\s+fuente|por\s+renta|de\s+renta)"
    r"|retenci[oó]n\s+en\s+la\s+fuente",
    re.I,
)

# Palabras que NUNCA son ciudades colombianas
_CIUDAD_BLACKLIST = frozenset({
    "day", "month", "year", "fecha", "periodo", "período", "avable", "gravable",
    "razon", "razón", "social", "nombre", "administración", "administracion",
    "impuestos", "certifica", "certificado", "retencion", "retención",
    "fuente", "durante", "contribuyente", "agente", "retenedor",
    "expedicion", "expedición", "consignacion", "consignación",
    "enero", "febrero", "marzo", "abril", "mayo", "junio",
    "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
    "año", "practicó", "practico", "practica", "siguiente", "siguientes",
    "valores", "empresa", "señores", "estimados", "retenido",
})

# Ciudad donde se practicó la retención (acepta variantes con/sin acento)
_RE_CIUDAD_PRACT = re.compile(
    r"ciudad\s+donde\s+se\s+[^\n:]{0,40}:([^.\n]{3,40})",
    re.I,
)
_RE_CIUDAD_CONSIG = re.compile(
    r"(?:"
    r"consignad[ao]s?\s+en\s+(?:la\s+(?:ciudad|municipio)\s+(?:de\s+)?)?"
    r"|retenciones?\s+consignadas?\s+en\s+(?:la\s+ciudad\s+(?:de\s+)?)?"
    r"|ciudad\s+donde\s+se\s+consig[^\n:]{0,25}[:\s]+"
    r")[:\s]*([A-ZÁÉÍÓÚÑ][^.\n]{2,35})",
    re.I,
)
_RE_CIUDAD_EMISOR = re.compile(r"^ciudad\s*:\s*([A-ZÁÉÍÓÚÑ][^,.\n]{2,25})", re.M | re.I)
# "EN LA CIUDAD DE MEDELLÍN" (formato narrativo)
_RE_CIUDAD_EN = re.compile(
    r"en\s+la\s+ciudad\s+de\s+([A-ZÁÉÍÓÚÑ][A-Za-záéíóúñ]{2,20})", re.I
)
# "Lugar de expedición: CIUDAD"
_RE_CIUDAD_LUGAR = re.compile(
    r"(?:lugar\s+de\s+expedi[cs]i[oó]n|expedido\s+en)\s*[:\-]?\s*([A-ZÁÉÍÓÚÑ][^\n,]{2,25})", re.I
)
_MESES_PAT = (
    r"(?:enero|febrero|marzo|abril|mayo|junio|julio|agosto"
    r"|septiembre|octubre|noviembre|diciembre)"
)
# "Medellín, 12 de enero de 2025" o "MEDELLÍN, 12 ENERO 2025"
_RE_CIUDAD_FECHA = re.compile(
    r"^([A-ZÁÉÍÓÚÑ][A-Za-záéíóúüñÑÁÉÍÓÚÜ]+(?:\s+[A-ZÁÉÍÓÚÜÑ][A-Za-záéíóúüñÑÁÉÍÓÚÜ]+)?)"
    r"\s*,\s*\d{1,2}\s+(?:de\s+)?" + _MESES_PAT,
    re.M | re.I,
)
# "Medellín, Enero 12 de 2025"
_RE_CIUDAD_FECHA2 = re.compile(
    r"^([A-ZÁÉÍÓÚÑ][A-Za-záéíóúüñÑÁÉÍÓÚÜ]+(?:\s+[A-ZÁÉÍÓÚÜÑ][A-Za-záéíóúüñÑÁÉÍÓÚÜ]+)?)"
    r"\s*,\s*" + _MESES_PAT + r"\s+\d{1,2}",
    re.M | re.I,
)
# "practicada/o en Medellín" o "practicada en la ciudad de Medellín"
_RE_CIUDAD_PRACTICADA = re.compile(
    r"practicad[ao]\s+en\s+(?:(?:la\s+ciudad|el\s+municipio)\s+de\s+)?"
    r"([A-ZÁÉÍÓÚÑ][A-Za-záéíóúüñÑÁÉÍÓÚÜ]+(?:\s+[A-ZÁÉÍÓÚÜÑ][A-Za-záéíóúüñÑÁÉÍÓÚÜ]+)?)",
    re.I,
)
# "municipio de Sabaneta" o "en el municipio de Itagüí"
_RE_CIUDAD_MUNICIPIO = re.compile(
    r"municipio\s+de\s+([A-ZÁÉÍÓÚÑ][A-Za-záéíóúüñÑÁÉÍÓÚÜ]+(?:\s+[A-ZÁÉÍÓÚÜÑ][A-Za-záéíóúüñÑÁÉÍÓÚÜ]+)?)",
    re.I,
)

# Línea TOTAL / TOTALES — hasta 3 columnas (base | pct% opcional | retención)
_RE_TOTAL_LINE = re.compile(
    r"^TOTAL(?:ES|S)?\s*(?:[A-Za-záéíóúñ\s]{0,25}?)\s*[:\-]?\s*"
    r"[\-\$\s]*([\d.,]+)\s+[\-\$\s]*([\d.,]+)(?:\s+[\-\$\s]*([\d.,]+))?",
    re.M | re.I,
)

# Tabla formato colombiano: Concepto | Base(X.XXX.XXX) | % | Ret(X.XXX)
_RE_TABLE_COL_FORMAT = re.compile(
    r"([A-Za-záéíóú][^\t\n]{2,40}?)\s+"
    r"(\d{1,3}(?:\.\d{3})+)\s+"      # base con puntos miles: 3.622.900
    r"[\d.,]+\s*%\s+"                 # porcentaje
    r"(\d{1,3}(?:\.\d{3})*(?:,\d+)?)",  # retención
    re.M,
)

# Tabla genérica: Concepto | Tasa% | Base | Retenido  (orden variado)
_RE_TABLE_TASA_BASE_RET = re.compile(
    r"([A-Za-záéíóú][^\t\n]{2,60}?)\s+([\d]+[,.]\d{1,4}\s*%?)\s+([\d.,]+)\s+([\d.,]+)",
    re.M,
)

# Formato "Retefuente 2.5%  BASE  TARIFA  VALOR_TOTAL" (base antes que tasa)
_RE_RETEFTE_FORMAT = re.compile(
    r"retefuente\s+[\d.]+%\s+([\d,.]+\.?\d*)\s+[\d,.]+\s*%\s+([\d,.]+)",
    re.I,
)

# Formato QUIRUSTETIC: "PERIODO CONCEPTO VALOR_BASE VALOR_RETENIDO" sin columna %
_RE_QUIRUSTETIC_ROW = re.compile(
    r"(?:Anual|Enero|Febrero|[A-Za-z]+)\s+(?:COMPRAS|SERVICIOS|VENTAS)[^\n]{0,30}\s+([\d,]+\.\d{2})\s+([\d,]+\.\d{2})",
    re.I,
)

# Formato SAN JUAN DE DIOS: código + concepto + % + "2,500" (código) + base + ret
_RE_SJD_ROW = re.compile(
    r"\d{8,}\s+[A-Z][^\n]{3,40}?\s+[\d.]+\s*%\s+[\d,]+\s+([\d.,]+)\s+([\d.,]+)",
    re.M,
)

# Párrafo MAYORISTA: "($ ret) m/l ... al pct% de ($ base)"
_RE_PARRAFO_MAYOR = re.compile(
    r"\(\$\s*([\d.,]+)\)\s*m/?l[^\n]{0,100}?(\d+[.,]\d+)\s*%[^\n]{0,60}?\(\$\s*([\d.,]+)\)",
    re.I,
)
# Párrafo MAYORISTA alternativo: "suma de $ret... equivalente al pct% de $base"
_RE_PARRAFO_MAYOR2 = re.compile(
    r"suma\s+de\s+(?:\(?\$\s*)?([\d.,]+)[^.]{0,100}?(\d+[.,]\d+)\s*%[^.]{0,80}?(?:\(?\$\s*)?([\d.,]+)",
    re.I,
)

# Tabla bimestre IVA (PROTECTION CLOUD): Bimestre | Base | IVA | Cuantía | %
_RE_BIMESTRE = re.compile(
    r"^\s*(?:[1-6]|[IVX]+)\s+\S[^\n]{5,80}?([\d,.]+)\s+([\d,.]+)\s+([\d,.]+)\s+([\d,.]+)",
    re.M,
)
# Tabla bimestre ROSARIO: I IVA DEL 19% Ret... | Base | IVA | % | Cuantia
_RE_BIMESTRE_ROSARIO = re.compile(
    r"^\s*(?:[IVX]+|[1-6])\s+IVA[^\n]{0,50}?([\d.,]+)\s+([\d.,]+)\s+[\d.,]+\s+([\d.,]+)",
    re.M,
)

# Tabla bimestre ICA mes-a-mes: "Enero - Febrero  base  ret" (IND FANTASIA ICA)
_RE_BIMESTRE_ICA = re.compile(
    r"^[A-Za-záéíóúñÑ]+\s*[-–]\s*[A-Za-záéíóúñÑ]+\s+([\d,.]+)\s+([\d,.]+)",
    re.M,
)
# Línea total ICA: "MONTO DEL PAGO SUJETO...\n base  ret" (IND FANTASIA ICA)
_RE_ICA_MONTO_TOTAL = re.compile(
    r"MONTO\s+DEL\s+PAGO[^\n]*\n\s*([\d,.]+)\s+([\d,.]+)",
    re.I,
)

# SUINSUMOYA box format
_RE_BOX_BASE = re.compile(
    r"RETENCION[^|$\n]{0,60}([\d.,]+)[^\d\n]*?([\d.,]+)",
    re.I,
)

# Formato PERSEO/etiqueta: "MONTO TOTAL: $X" + "CUANTIA DE LA RETENCION: $Y"
_RE_MONTO_TOTAL = re.compile(
    r"MONTO\s+TOTAL[:\s]*\$?\s*([\d.,]+)", re.I
)
_RE_CUANTIA = re.compile(
    r"(?:CUANTIA(?:\s+DE\s+LA\s+RETENCION)?|VALOR\s+RETENID[OA]|VALOR\s+TOTAL)[:\s]*\$?\s*([\d.,]+)", re.I
)
# Formato minimalista: "BASE GRAVABLE: $X" / "RETENCIÓN EN LA FUENTE: $Y"
_RE_BASE_LABEL = re.compile(
    r"(?:BASE\s+(?:GRAVABLE|SUJETA(?:\s+A\s+RETENCI[OÓ]N)?|DE\s+RETENCI[OÓ]N|IMPONIBLE)?|VALOR\s+BASE|MONTO\s+BASE)[:\s]*\$?\s*([\d.,]+)",
    re.I,
)
_RE_RET_LABEL = re.compile(
    r"(?:RETENCI[OÓ]N(?:\s+EN\s+LA\s+FUENTE)?|RETEFUENTE|RTE\.?\s*(?:FTE\.?|FUENTE|RENTA)?|VALOR\s+RETENI[DO]|CUANT[IÍ]A(?:\s+RETENI[DA])?)[:\s]*\$?\s*([\d.,]+)",
    re.I,
)


# ── Patrones adicionales de NIT por layout ───────────────────────────────────

# "Retenedor : 900742492-9" (Mekano ERP — ENTREAGUAS)
_RE_RETENEDOR_NIT = re.compile(
    r"[Rr]etenedor\s*:\s*([0-9]{6,12}|[0-9]{1,3}(?:[.,][0-9]{3})+)\s*[-–]\s*([0-9])\b",
)
# "Agente retenedor: NOMBRE" o "Agente Retenedor\t\tNOMBRE" (COMERCIALIZADORA CSF / XLS)
_RE_AGENTE_RETENEDOR_LABEL = re.compile(
    r"[Aa]gente\s+[Rr]etenedor\s*(?::\s*|[\t ]{2,})([^\n\t]{3,80})", re.I
)
# "NIT o Cédula" o "Nit ó Cédula" (con o sin acento) — también tab-separado (XLS)
_RE_NIT_O_CEDULA = re.compile(
    r"NIT\s+[oó]\s+C[eé]dula\s*[:\t\s]+([0-9]{6,12}|[0-9]{1,3}(?:[.,][0-9]{3})+)\s*[-–]\s*([0-9])\b",
    re.I,
)
_RE_NIT_O_CEDULA_NODV = re.compile(
    r"NIT\s+[oó]\s+C[eé]dula\s*[:\t\s]+([0-9]{6,12}|[0-9]{1,3}(?:[.,][0-9]{3})+)",
    re.I,
)
# "NIT. 900,713,526 - 7"  (PUBLIK MAGIC — miles con coma)
_RE_NIT_COMMA_MILES = re.compile(
    r"N\.?I\.?T\.?\s*[.:\s#]*([0-9]{1,3}(?:,[0-9]{3})+)\s*[-–]\s*([0-9])\b",
    re.I,
)
# NIT bare con puntos en línea sola: "890.906.574-4" (COLEGIO ALEMAN)
_RE_BARE_NIT_DOTS_DV = re.compile(r"^([0-9]{1,3}(?:\.[0-9]{3})+)[-–]([0-9])$")
# SAP bilingüe: "Company Code Tax No. 8001974634" (EL BUCANERO)
_RE_SAP_COMPANY_NIT = re.compile(r"Company\s+Code\s+Tax\s+No\.?\s+([0-9]{7,12})", re.I)
_RE_SAP_COMPANY_NAME = re.compile(r"Razon\s+social\s+quien[^\n]*\n([^\n]+)", re.I)
# "Razón Social del Agente Retenedor ... Nit o C.C." seguido en línea siguiente del dato
_RE_AGENTE_RETENEDOR_HEADER = re.compile(
    r"Raz[oó]n\s+[Ss]ocial[^\n]{0,60}?[Aa]gente\s+[Rr]etenedor[^\n]*\n([^\n]+)", re.I
)
# Párrafo narrativo: "suma de $RET ... base de $BASE" — [^$] permite saltos de línea
_RE_NARRATIVO_BASE = re.compile(
    r"suma\s+de\s+\$\s*([\d.,]+)[^$]{0,400}?base\s+de\s+\$\s*([\d.,]+)",
    re.I | re.DOTALL,
)
# Formato MEDIFE/IRCC: "COD8DIG  Concepto  PCT%  PCT2  $BASE  $RET"
_RE_MEDIFE_ROW = re.compile(
    r"^\d{5,10}\s+[A-Za-záéíóú][^\n]{2,40}?[\d.,]+\s*%\s+[\d.,]+\s+\$\s*([\d.,]+)\s+\$\s*([\d.,]+)",
    re.M,
)
# Montos entre paréntesis: "($9,396,720) ($234,918)" (PUBLIK MAGIC)
_RE_PAREN_AMOUNTS = re.compile(
    r"\(\$?\s*([\d.,]{4,})\)\s+\(\$?\s*([\d.,]{3,})\)",
    re.I,
)
# SAP total line: "(TOTAL)... base   pct   ret"
_RE_SAP_TOTAL = re.compile(
    r"\(TOTAL\)[^\n]{0,80}?([\d,]{5,})\s+[\d.,]+\s+([\d,]{3,})\s*$",
    re.M | re.I,
)

# ── NIT y Razón Social ────────────────────────────────────────────────────────

# Patrones de texto que NO son razón social
_RE_RAZON_INVALID = re.compile(
    r"^(?:fecha\s+de\s+expedic|año\s+gravable|periodo\s+gravable"
    r"|certificado\s+de\s+retenci|periodo\s+de\s+retenci)",
    re.I,
)


def _clean_nit(raw: str) -> str:
    """Elimina puntos, comas y guiones de un NIT: '800.233.836' → '800233836'."""
    return re.sub(r"\D", "", raw)


_RE_BARE_NIT_DV  = re.compile(r"^(\d{7,12})\s*[-–]\s*(\d)$")
_RE_BARE_NIT     = re.compile(r"^(\d{7,12})[-–]?$")   # acepta trailing dash (890905154-)
_RE_RAZON_LABEL  = re.compile(r"razón?\s+social[^:]{0,20}:\s*([^\n]{3,70})", re.I)
# Líneas de paginación a ignorar al buscar razón social
_RE_PAGINA       = re.compile(r"^p[aá]g(?:ina)?[.:]?\s*\d+$", re.I)
# Líneas que parecen dirección (no son razón social)
_RE_ADDR_START   = re.compile(
    r"^(?:CL|CR|CRA|KR|AK|AV|KM|VDA|CARRERA|CALLE|DIAGONAL|TRANSVERSAL|AVENIDA|BARRIO|SECTOR)\b",
    re.I,
)


def _is_junk_line(ln: str) -> bool:
    """Retorna True si la línea es paginación, dirección, teléfono o demasiado corta para ser razón social."""
    if not ln or len(ln) < 3:
        return True
    if _RE_PAGINA.match(ln):
        return True
    if _RE_ADDR_START.match(ln):
        return True
    # Línea que es sólo caracteres de guión/relleno
    if re.match(r"^[-=_*·.]{4,}$", ln):
        return True
    # Línea de solo dígitos (teléfono, código, etc.) — no es una razón social
    if re.match(r"^[\d\s\-+().]{6,}$", ln) and re.search(r"\d{6,}", ln):
        return True
    # Líneas de título de documento (no son nombres de empresa)
    if re.match(r"^(?:certificado|periodo\s+gravable|a[ñn]o\s+gravable|fecha\s+de\s+expedi|constancia\s+de|datos\s+del|certifica\s+que)", ln, re.I):
        return True
    return False


def _extract_emisor(text: str) -> tuple[str, str, str]:
    """
    Retorna (razon_social, nit_limpio, dv) del AGENTE RETENEDOR (quien retiene).
    Usa detección por prioridad para manejar múltiples layouts.
    """

    def _make_result(razon: str, nit_raw: str, dv_raw: str) -> tuple[str, str, str]:
        nit = _clean_nit(nit_raw)
        d = dv_raw if dv_raw else (calcular_dv(nit) if nit else "")
        # Strip label prefixes (Retenedor:, Señores:, RETENIDO:, etc.)
        razon = re.sub(
            r"(?i)^(?:retenedor|retenido|señores?|señor|estimados?)\s*:\s*",
            "", razon,
        ).strip()
        razon = re.sub(r"(?i)(agente\s+retenedor|razón?\s+social|nombre\s+comercial|empresa)[:\s]*", "", razon).strip()
        razon = re.sub(r"[\|_\-]{2,}", "", razon).strip()
        # Strip leading NIT prefix (e.g. "901996294-SIEMPRE NIÑAS")
        razon = re.sub(r"^\d{6,12}[-–]\s*", "", razon).strip()
        # Discard header lines that are not a company name
        if _RE_RAZON_INVALID.search(razon):
            razon = ""
        return (razon if len(razon) >= 3 else ""), nit, d

    # ── PRIORIDAD 1: SAP bilingüe (varias variantes de etiqueta) (EL BUCANERO) ──
    # El NIT aparece al FINAL de la línea siguiente al header de retenedor
    for _sap_pat in (
        r"Company\s+Code\s+Tax\s+No\.?[^\n]*\n([^\n]+)",
        r"Company\s+Code\s+Name\s+NIT[^\n]*\n([^\n]+)",
    ):
        m = re.search(_sap_pat, text, re.I)
        if m:
            next_line = m.group(1).strip()
            # NIT al final de la línea (layout RENTA)
            nit_m = re.search(r"(\d{7,12})\s*$", next_line)
            if nit_m:
                nit_raw = nit_m.group(1)
                razon = next_line[:nit_m.start()].strip()
                return _make_result(razon, nit_raw, "")
            # NIT en cualquier posición de la línea (layout IVA u otras variantes)
            nit_m = re.search(r"(\d{7,12})", next_line)
            if nit_m:
                nit_raw = nit_m.group(1)
                razon = next_line[:nit_m.start()].strip()
                if not razon:
                    # Buscar nombre en líneas anteriores al header SAP
                    prev_lines = [l.strip() for l in text[:m.start()].split("\n") if l.strip()]
                    for ln in reversed(prev_lines[-5:]):
                        if not _is_junk_line(ln):
                            razon = ln
                            break
                return _make_result(razon, nit_raw, "")
            # Header SAP encontrado pero sin NIT reconocible: NO hacer break,
            # dejar que las prioridades siguientes intenten extraer
    # (no break aquí — continuar con prioridades 2..5 si SAP no sirvió)

    # ── PRIORIDAD 2: "Razón Social del Agente Retenedor" sección (MEDIFE) ──────
    m = _RE_AGENTE_RETENEDOR_HEADER.search(text)
    if m:
        row_line = m.group(1).strip()
        # La línea tiene: RazonSocial  Dirección  NIT (sin guión, al final)
        nit_m = re.search(r"(\d{6,12})\s*[-–]?\s*(\d?)\s*$", row_line)
        if nit_m:
            nit_raw = nit_m.group(1)
            dv = nit_m.group(2) or ""
            razon = row_line[:nit_m.start()].strip()
            # Quitar dirección (empieza con Cra, Cl, Av, etc.)
            razon = re.split(r"\s{2,}|(?:Cra|Cll|Cl\b|Cr\b|AV|KM|CR)\s", razon, flags=re.I)[0].strip()
            return _make_result(razon, nit_raw, dv)

    # ── PRIORIDAD 3: "Agente retenedor: NOMBRE" + NIT en línea siguiente ────────
    m = _RE_AGENTE_RETENEDOR_LABEL.search(text)
    if m:
        razon = m.group(1).strip()
        snippet = text[m.end():m.end() + 400]
        nm = (_RE_NIT_O_CEDULA.search(snippet)
              or _RE_NIT_COMMA_MILES.search(snippet)
              or _RE_NIT_DV.search(snippet)
              or _RE_NIT_O_CEDULA_NODV.search(snippet)
              or _RE_NIT_NODV.search(snippet))
        if nm:
            nit_raw = nm.group(1)
            dv = nm.group(2) if nm.lastindex and nm.lastindex >= 2 else ""
            return _make_result(razon, nit_raw, dv)

    # ── PRIORIDAD 4b: tabla SAP "Nombre EMPRESA Nit: NNN - D" en una sola línea ──
    # [^\S\n]+ en vez de \s+ evita cruzar líneas y capturar el "Retenido a"
    m = re.search(
        r"[Nn]ombre[^\S\n]+([A-ZÁÉÍÓÚÑ&][^\n]{3,60}?)[^\S\n]+[Nn]it[:\s]+([0-9]{1,3}(?:[.,][0-9]{3})+|[0-9]{7,12})\s*[-–]\s*([0-9])",
        text, re.I,
    )
    if m:
        razon_cand = m.group(1).strip()
        nit_cand   = re.sub(r"[.,]", "", m.group(2))
        dv_cand    = m.group(3)
        if len(nit_cand) >= 7:
            return _make_result(razon_cand, nit_cand, dv_cand)

    # ── PRIORIDAD 4: "Retenedor : NIT-DV" (Mekano ERP — ENTREAGUAS) ─────────────
    m = _RE_RETENEDOR_NIT.search(text)
    if m:
        nit_raw, dv = m.group(1), m.group(2)
        prev = text[:m.start()]
        prev_lines = [l.strip() for l in prev.split("\n") if l.strip()]
        razon = ""
        for ln in reversed(prev_lines[-5:]):
            if _is_junk_line(ln):
                continue
            # "Agente NOMBRE" → quitar prefijo "Agente"
            razon = re.sub(r"^[Aa]gente\s+", "", ln).strip()
            break
        return _make_result(razon, nit_raw, dv)

    # ── PRIORIDAD 5: lógica original (bare NIT + etiqueta NIT) ──────────────────
    lines = [ln.strip() for ln in text.split("\n") if ln.strip()]
    nit_raw = dv = ""
    nit_line_idx = -1

    for i, ln in enumerate(lines[:8]):
        m2 = _RE_BARE_NIT_DOTS_DV.match(ln)
        if m2:
            nit_raw, dv = m2.group(1), m2.group(2)
            nit_line_idx = i
            break
        m2 = _RE_BARE_NIT_DV.match(ln)
        if m2:
            nit_raw, dv = m2.group(1), m2.group(2)
            nit_line_idx = i
            break
        m2 = _RE_BARE_NIT.match(ln)
        if m2 and len(m2.group(1)) >= 8:
            candidate = re.sub(r"[\.,]", "", m2.group(1))
            # Saltar números de 10 dígitos que empiezan con 3 (celulares colombianos)
            if len(candidate) == 10 and candidate.startswith("3"):
                continue
            nit_raw = m2.group(1)
            nit_line_idx = i
            break

    if not nit_raw:
        for i, ln in enumerate(lines[:30]):
            m2 = (_RE_NIT_COMMA_MILES.search(ln)
                  or _RE_NIT_DV.search(ln)
                  or _RE_NIT_NODV.search(ln))
            if m2:
                nit_raw = m2.group(1)
                dv = m2.group(2) if m2.lastindex and m2.lastindex >= 2 else ""
                nit_line_idx = i
                break

    nit_clean = _clean_nit(nit_raw)
    if not dv and nit_clean:
        dv = calcular_dv(nit_clean)

    razon = ""
    if nit_line_idx > 0:
        city_fallback = ""
        for look_back in range(1, min(nit_line_idx + 1, 6)):
            candidate = lines[nit_line_idx - look_back]
            if _RE_NIT_DV.search(candidate) or _RE_NIT_NODV.search(candidate):
                razon = re.split(r"N\.?I\.?T", candidate, flags=re.I)[0].strip()
                break
            if _is_junk_line(candidate):
                continue
            # Saltar línea de una sola palabra en mayúsculas sin sufijo empresarial
            # (probable ciudad o departamento), pero guardar como fallback
            words = candidate.strip().split()
            is_single_caps = (
                len(words) == 1
                and re.match(r"^[A-ZÁÉÍÓÚÑ]{3,15}$", words[0])
                and not re.search(r"\b(?:SAS|LTDA|S\.A\.S|S\.A\.|E\.U\.|CORP)\b", candidate, re.I)
            )
            if is_single_caps:
                if not city_fallback:
                    city_fallback = candidate
                continue
            # Verificar si el candidato multipalabra es un municipio conocido
            cod_dpto, _ = buscar_municipio(candidate.strip())
            if cod_dpto:
                if not city_fallback:
                    city_fallback = candidate
                continue
            razon = candidate
            break
        if not razon and city_fallback:
            # Usar city_fallback solo si NO es un municipio conocido
            cod_dpto, _ = buscar_municipio(city_fallback.strip())
            if not cod_dpto:
                razon = city_fallback
    elif nit_line_idx == 0 and len(lines) > 1:
        razon = lines[1] if not _RE_NIT_NODV.search(lines[1]) else ""

    if len(razon) < 3:
        m2 = _RE_RAZON_LABEL.search(text[:800])
        if m2:
            razon = m2.group(1).strip()

    return _make_result(razon, nit_raw if nit_raw else "", dv)


def _extract_direccion(text: str) -> str:
    m = re.search(
        r"(?:(?<!\w)(?:CLL?|CRA?|KR|AK|AV|KM|VDA)(?!\w)|CARRERA|CALLE|DIAGONAL|TRANSVERSAL|AVENIDA)[^\n]{3,70}",
        text, re.I,
    )
    if not m:
        return ""
    addr = m.group(0).strip()
    # Quitar teléfonos y texto de relleno al final
    addr = re.sub(r"\s+(?:Tel[eé]fono|Telf?|Cel|Fax|PBX)\.?\s*[\d\s()\-+.]{6,}.*$", "", addr, flags=re.I)
    addr = re.sub(r"\s+[-–]\s*(?:Tel[eé]fono|Tel|Cel).*$", "", addr, flags=re.I)
    addr = re.sub(r"\s+\+?57[\s\-]?\d{7,}.*$", "", addr)  # número colombiano directo
    addr = re.sub(r"\s+\d{7,}.*$", "", addr)               # teléfono sin prefijo
    return addr.strip()[:80]


# ── Persona natural — separar apellidos y nombres ────────────────────────────

def _split_persona_natural(full_name: str) -> tuple[str, str, str, str]:
    """
    Para tipo_doc='13', intenta separar apellidos y nombres.
    Convención DIAN: Apellido1 Apellido2 Nombre1 OtrosNombres.
    Heurística: si >= 4 tokens → últimos 2 = apellidos, primeros = nombres.
    """
    # Si hay "/" separando nombre comercial del nombre real, tomar la primera parte
    clean = re.split(r"\s*/\s*", full_name)[0].strip()
    # Quitar sufijos como "E.U.", "S.A.S." (no son nombres de persona)
    clean = re.sub(r"\b(?:E\.U\.|S\.A\.S\.?|S\.A\.|LTDA\.?|S\.C\.S\.?)\s*$", "", clean, flags=re.I).strip()
    tokens = clean.split()
    if not tokens:
        return "", "", "", ""
    if len(tokens) == 1:
        return tokens[0], "", "", ""
    if len(tokens) == 2:
        # Podría ser "NombreApellido" o "ApellidoNombre" — asumimos Apellido Nombre
        return tokens[0], "", tokens[1], ""
    if len(tokens) == 3:
        return tokens[0], tokens[1], tokens[2], ""
    # 4+ tokens: últimos 2 = apellidos, primeros = nombres (convención DIAN)
    return tokens[-2], tokens[-1], tokens[0], " ".join(tokens[1:-2])


# ── Tipo de certificado ───────────────────────────────────────────────────────

def _cert_type(text: str) -> str:
    """Devuelve 'ICA', 'IVA' o 'RENTA'."""
    snippet = text[:800]
    if _RE_CERT_ICA.search(snippet):
        return "ICA"
    if _RE_CERT_IVA.search(snippet):
        return "IVA"
    if _RE_CERT_RENTA.search(snippet):
        return "RENTA"
    # Fallback: si menciona "fuente" o "renta" → RENTA
    if re.search(r"\b(?:fuente|renta)\b", snippet, re.I):
        return "RENTA"
    return "RENTA"


# ── Concepto 1003 — todos los conceptos DIAN Formato 1003 ───────────────────

# Mapa ordenado por especificidad (más específico primero)
_CONCEPTO_MAP = [
    # 1301: Salarios y pagos laborales
    (re.compile(r"SALARIOS?|SUELDOS?|PAGOS?\s+LABORALES?|LIQUIDACI[OÓ]N\s+LABORAL|NOMINA", re.I), "1301"),
    # 1304: Honorarios (antes de servicios para evitar solapamiento)
    (re.compile(r"HONORARIOS?", re.I), "1304"),
    # 1305: Comisiones
    (re.compile(r"COMISIONES?", re.I), "1305"),
    # 1307: Arrendamientos
    (re.compile(r"ARRENDAMIENTOS?|RENTA\s+DE\s+BIEN(?:ES)?|CANON\s+DE\s+ARRENDAMIENTO", re.I), "1307"),
    # 1306: Intereses y rendimientos financieros
    (re.compile(r"INTERESES?|RENDIMIENTOS?\s+FINANCIEROS?|RENDIMIENTO\s+FINANCIERO", re.I), "1306"),
    # 1308: Regalías
    (re.compile(r"REGAL[IÍ]AS?|PROPIEDAD\s+INTELECTUAL|DERECHOS?\s+DE\s+AUTOR", re.I), "1308"),
    # 1310: Dividendos y participaciones
    (re.compile(r"DIVIDENDOS?|PARTICIPACIONES?", re.I), "1310"),
    # 1311: Enajenación de activos fijos
    (re.compile(r"ENAJENACI[OÓ]N|ACTIVOS?\s+FIJOS?|VENTA\s+DE\s+ACTIVOS?", re.I), "1311"),
    # 1312: Pagos tarjetas débito/crédito
    (re.compile(r"TARJETAS?\s+(?:D[EÉ]BITO|CR[EÉ]DITO)|INGRESOS?\s+DE\s+TARJETAS?", re.I), "1312"),
    # 1313: Loterías, rifas, apuestas
    (re.compile(r"LOTER[IÍ]AS?|RIFAS?|APUESTAS?|JUEGOS?\s+DE\s+AZAR", re.I), "1313"),
    # 1314: Contratos de construcción
    (re.compile(r"CONSTRUCCI[OÓ]N|URBANIZACI[OÓ]N|OBRA\s+(?:CIVIL|PUBLICA)", re.I), "1314"),
    # 1315: Pagos al exterior
    (re.compile(r"(?:PAGOS?\s+)?AL\s+EXTERIOR|REMESAS?|GIROS?\s+AL\s+EXTERIOR", re.I), "1315"),
    # 1316: Títulos de renta fija
    (re.compile(r"T[IÍ]TULOS?\s+(?:DE\s+)?RENTA\s+FIJA|BONOS?\s+DE\s+DEUDA", re.I), "1316"),
    # 1317: Trabajadores independientes
    (re.compile(r"TRABAJADORES?\s+INDEPENDIENTES?|CONTRATISTAS?\s+INDEPENDIENTES?", re.I), "1317"),
    # 1303: Servicios generales (después de los más específicos)
    (re.compile(r"SERVICIOS?\b", re.I), "1303"),
    # SAP bilingüe — W-INC Purchase → 1302, W-INC Services → 1303
    (re.compile(r"W-INC\s+Purch|Purch\.other\s+income", re.I), "1302"),
    (re.compile(r"W-INC\s+Serv", re.I), "1303"),
    # 1302: Compras (al final como más general)
    (re.compile(r"COMPRAS?\b|COMPRA\s+GENERAL|COMPRAS?\s+GENERALES?", re.I), "1302"),
    (re.compile(r"VENTAS?\b", re.I), "1302"),
]


def _pct_to_concepto(keyword: str, pct: float) -> str:
    """Mapea (palabra del concepto, porcentaje) → código DIAN."""
    for pattern, code in _CONCEPTO_MAP:
        if pattern.search(keyword):
            return code
    # Fallback por tasa: >=3.5% generalmente servicios/honorarios
    if pct >= 3.5:
        return "1303"
    return "1302"


# Porcentajes por defecto si no se detecta tasa en el texto
_PCT_DEFAULT = {
    "1301": 0.0, "1302": 2.5, "1303": 4.0, "1304": 10.0, "1305": 11.0,
    "1306": 7.0, "1307": 3.5, "1308": 3.5, "1309": 15.0, "1310": 0.0,
    "1311": 1.0, "1312": 1.5, "1313": 20.0, "1314": 2.0, "1315": 0.0,
    "1316": 4.0, "1317": 2.0, "ICA": 0.0,
}


def _detect_concepto(text: str, tipo_cert: str) -> tuple[str, float]:
    """Retorna (concepto_1003, porcentaje)."""
    if tipo_cert == "ICA":
        return "ICA", 0.0
    if tipo_cert == "IVA":
        return "1309", 15.0

    pct_m = re.search(r"(\d+[.,]\d+)\s*%", text)
    pct = float(pct_m.group(1).replace(",", ".")) if pct_m else 0.0

    for pattern, code in _CONCEPTO_MAP:
        if pattern.search(text):
            return code, pct if pct else _PCT_DEFAULT.get(code, 2.5)

    return "1302", pct if pct else 2.5


# ── Extracción de filas múltiples de conceptos ────────────────────────────────

# Tabla "Concepto LABEL PCT%  BASE  RETENCIÓN" — múltiples filas
# (?!\s*%) evita capturar valores como "2.500%" (porcentajes repetidos) como ret
_RE_MULTI_ROW = re.compile(
    r"(?P<label>[A-Za-záéíóúÁÉÍÓÚñÑ][^\n\t]{2,50}?)\s+"
    r"(?P<pct>[\d]+[.,][\d]{1,4})\s*%\s+"
    r"(?P<base>[\d.,]{4,}(?:\.\d{2})?)\s+%?\s*"
    r"(?P<ret>[\d.,]{3,})(?!\s*%)",
    re.M,
)

# Tabla "Concepto (PCT%)  num1  num2  num3  ..." — pct entre paréntesis (GASCAL)
_RE_MULTI_ROW_PAREN_PCT = re.compile(
    r"(?P<label>[A-Za-záéíóúÁÉÍÓÚñÑ][^\n\t(]{2,40}?)\s*\((?P<pct>[\d]+[.,][\d]{1,4})\s*%\)\s+"
    r"(?P<rest>[^\n]+)",
    re.M,
)

# Tabla sin porcentaje explícito: "CONCEPTO  BASE  RETENCIÓN"
_RE_MULTI_ROW_NOPCT = re.compile(
    r"(?P<label>COMPRAS?|SERVICIOS?|HONORARIOS?|ARRENDAMIENTOS?|COMISIONES?|VENTAS?)[^\n]{0,40}\s+"
    r"(?P<base>[\d.,]{4,}(?:\.\d{2})?)\s+"
    r"(?P<ret>[\d.,]{3,})",
    re.M | re.I,
)


def _extract_concepto_rows(
    text: str, tipo_cert: str
) -> list[tuple[str, float, float, float]]:
    """
    Intenta encontrar todas las filas de conceptos en el certificado.
    Retorna lista de (concepto_dian, porcentaje, base, retencion).
    Si no encuentra filas múltiples, retorna lista vacía.
    """
    if tipo_cert in ("ICA", "IVA"):
        return []

    rows: list[tuple[str, float, float, float]] = []

    # 1. Tabla con porcentaje explícito
    for m in _RE_MULTI_ROW.finditer(text):
        label = m.group("label").strip()
        pct   = float(m.group("pct").replace(",", "."))
        base  = _parse_money(m.group("base"))
        ret   = _parse_money(m.group("ret"))
        if base > 100 and ret > 0 and base > ret:
            if 0.001 < ret / base < 0.35:
                concepto = _pct_to_concepto(label, pct)
                rows.append((concepto, pct, base, ret))
            elif pct > 0:
                # Ratio inválido pero pct conocido (tabla multi-columna tipo CUEROS VELEZ):
                # buscar el par correcto en el área del match usando pct como guía.
                pct_frac = pct / 100
                area = text[m.start("base"):m.start("base") + 300]
                area_nums = [_parse_money(n) for n in re.findall(r"[\d.,]+", area)]
                area_nums = [n for n in area_nums if 1_000 < n < 5_000_000_000]
                found = False
                for i in range(len(area_nums)):
                    for j in range(i + 1, min(i + 5, len(area_nums))):
                        b2, r2 = area_nums[i], area_nums[j]
                        if b2 > r2 > 100 and abs(r2 / b2 - pct_frac) < pct_frac * 0.12:
                            concepto = _pct_to_concepto(label, pct)
                            rows.append((concepto, pct, b2, r2))
                            found = True
                            break
                    if found:
                        break

    if rows:
        return rows

    # 1b. Tabla con pct entre paréntesis: "Concepto (PCT%) num1 num2 ..." (GASCAL)
    for m in _RE_MULTI_ROW_PAREN_PCT.finditer(text):
        label = m.group("label").strip()
        pct   = float(m.group("pct").replace(",", "."))
        rest  = m.group("rest")
        pct_frac = pct / 100
        nums_in_row = [_parse_money(n) for n in re.findall(r"[\d.,]+", rest)]
        nums_in_row = [n for n in nums_in_row if 1_000 < n < 5_000_000_000]
        for i in range(len(nums_in_row)):
            for j in range(i + 1, min(i + 5, len(nums_in_row))):
                b2, r2 = nums_in_row[i], nums_in_row[j]
                if b2 > r2 > 100 and abs(r2 / b2 - pct_frac) < pct_frac * 0.12:
                    concepto = _pct_to_concepto(label, pct)
                    rows.append((concepto, pct, b2, r2))
                    break
            if rows:
                break

    if rows:
        return rows

    # 2. Tabla sin porcentaje (compras/servicios sin tasa)
    for m in _RE_MULTI_ROW_NOPCT.finditer(text):
        label = m.group("label").strip()
        base  = _parse_money(m.group("base"))
        ret   = _parse_money(m.group("ret"))
        # Validación de ratio para evitar confundir montos y porcentajes
        if base > 500 and ret > 0 and base > ret and 0.003 < ret / base < 0.30:
            concepto = _pct_to_concepto(label, 0.0)
            pct = round(ret / base * 100, 2) if base else 0.0
            rows.append((concepto, pct, base, ret))

    return rows


# ── Extracción de montos ──────────────────────────────────────────────────────

def _extract_amounts(text: str, tipo_cert: str, nit_hint: str = "") -> tuple[float, float]:
    """Retorna (base, retencion). Intenta múltiples layouts."""

    # 1. Línea TOTAL/TOTALES — hasta 3 columnas: base | [pct%] | retención
    m = _RE_TOTAL_LINE.search(text)
    if m:
        n1 = _parse_money(m.group(1))
        n2 = _parse_money(m.group(2))
        n3 = _parse_money(m.group(3)) if m.group(3) else 0.0
        # Si hay 3 números y n2 parece porcentaje (ratio n2/n1 < 0.001 ó n2 ≤ 100),
        # usar (n1, n3) en lugar de (n1, n2)
        if n1 > 0 and n3 > 0 and (n2 <= 100 or n2 / n1 < 0.001) and 0.001 < n3 / n1 < 0.35:
            b, r = n1, n3
        else:
            b, r = n1, n2
        if b > 0 and r > 0:
            if tipo_cert == "ICA":
                # ICA: tasas 0.3%–2%, bases reales > 10.000 COP
                if b > 10_000 and r > 100 and 0.001 < r / b < 0.025:
                    return b, r
            elif 0.001 < r / b < 0.35:
                return b, r

    # 1b. ICA — sumar filas bimestrales "Mes - Mes  base  ret" (IND FANTASIA ICA)
    if tipo_cert == "ICA":
        m_ica = _RE_ICA_MONTO_TOTAL.search(text)
        if m_ica:
            b, r = _parse_money(m_ica.group(1)), _parse_money(m_ica.group(2))
            if b > r > 0:
                return b, r
        bims_ica = _RE_BIMESTRE_ICA.findall(text)
        if bims_ica:
            base_sum = sum(_parse_money(b) for b, _ in bims_ica)
            ret_sum  = sum(_parse_money(r) for _, r in bims_ica)
            if base_sum > 0:
                return base_sum, ret_sum

    # 1c. Formato PERSEO etiqueta: "MONTO TOTAL: $X  CUANTIA: $Y"
    m_total = _RE_MONTO_TOTAL.search(text)
    m_cuant  = _RE_CUANTIA.search(text)
    if m_total and m_cuant and m_total.start() < m_cuant.start():
        b = _parse_money(m_total.group(1))
        r = _parse_money(m_cuant.group(1))
        if b > r > 0:
            return b, r

    # 2. Formato Retefuente (base antes de tarifa)
    m = _RE_RETEFTE_FORMAT.search(text)
    if m:
        return _parse_money(m.group(1)), _parse_money(m.group(2))

    # 3. Tabla formato colombiano X.XXX.XXX (MONTERROZA style)
    rows_col = _RE_TABLE_COL_FORMAT.findall(text)
    if rows_col:
        base_sum = ret_sum = 0.0
        for _conc, base_s, ret_s in rows_col:
            b, r = _parse_money(base_s), _parse_money(ret_s)
            if b > 500 and r > 0 and b > r:
                base_sum += b
                ret_sum  += r
        if base_sum > 0:
            return base_sum, ret_sum

    # 4. Formato QUIRUSTETIC (Anual COMPRAS base ret sin %)
    m = _RE_QUIRUSTETIC_ROW.search(text)
    if m:
        return _parse_money(m.group(1)), _parse_money(m.group(2))

    # 4. Párrafo MAYORISTA con paréntesis
    m = _RE_PARRAFO_MAYOR.search(text)
    if m:
        return _parse_money(m.group(3)), _parse_money(m.group(1))

    m = _RE_PARRAFO_MAYOR2.search(text)
    if m:
        return _parse_money(m.group(3)), _parse_money(m.group(1))

    # 5. Tabla bimestre ROSARIO (IVA por bimestre con totales)
    ros = _RE_BIMESTRE_ROSARIO.findall(text)
    if ros:
        return sum(_parse_money(r[0]) for r in ros), sum(_parse_money(r[2]) for r in ros)

    # 6. Tabla bimestre genérica (PROTECTION CLOUD)
    bims = _RE_BIMESTRE.findall(text)
    if bims:
        # Columnas típicas: base, monto_iva, cuantia, porcentaje
        iva = tipo_cert == "IVA"
        base_sum = sum(_parse_money(r[0]) for r in bims)
        ret_sum  = sum(_parse_money(r[2]) for r in bims) if iva else sum(_parse_money(r[1]) for r in bims)
        if base_sum > 100:
            return base_sum, ret_sum

    # 7. Formato SAN JUAN DE DIOS (código + concepto + % + código + base + ret)
    m = _RE_SJD_ROW.search(text)
    if m:
        return _parse_money(m.group(1)), _parse_money(m.group(2))

    # 8. Tabla genérica 4 columnas (Concepto | Tasa% | Base | Ret)
    rows = _RE_TABLE_TASA_BASE_RET.findall(text)
    if rows:
        base_sum = ret_sum = 0.0
        for _conc, _tasa, base_s, ret_s in rows:
            b, r = _parse_money(base_s), _parse_money(ret_s)
            # Validar rango de tasa (0.1% a 30%) — evita capturar porcentaje como monto
            if b > 500 and r > 0 and b > r and 0.001 < r / b < 0.30:
                base_sum += b
                ret_sum  += r
        if base_sum > 0:
            return base_sum, ret_sum

    # 9. SUINSUMOYA box
    m = _RE_BOX_BASE.search(text)
    if m:
        b, r = _parse_money(m.group(1)), _parse_money(m.group(2))
        if b > r > 0:
            return b, r

    # 10. Formato MEDIFE/IRCC: código_concepto + label + % + pct2 + $BASE + $RET
    m = _RE_MEDIFE_ROW.search(text)
    if m:
        b, r = _parse_money(m.group(1)), _parse_money(m.group(2))
        if b > r > 0 and 0.003 < r / b < 0.30:
            return b, r

    # 11. Párrafo narrativo "La suma de $RET ... base de $BASE" (PROMOCIONES PUBLICITARIAS)
    m = _RE_NARRATIVO_BASE.search(text)
    if m:
        r, b = _parse_money(m.group(1)), _parse_money(m.group(2))
        if b > r > 0:
            return b, r

    # 12. Montos entre paréntesis "($9,396,720) ($234,918)" (PUBLIK MAGIC)
    m = _RE_PAREN_AMOUNTS.search(text)
    if m:
        b, r = _parse_money(m.group(1)), _parse_money(m.group(2))
        if b > r > 0 and 0.003 < r / b < 0.30:
            return b, r

    # 13. SAP total line "(TOTAL)... base  pct  ret" (EL BUCANERO)
    m = _RE_SAP_TOTAL.search(text)
    if m:
        b, r = _parse_money(m.group(1)), _parse_money(m.group(2))
        if b > r > 0:
            return b, r

    # 14. Formato minimalista etiquetado: "BASE GRAVABLE: X" + "RETENCIÓN: Y"
    m_base = _RE_BASE_LABEL.search(text)
    m_ret  = _RE_RET_LABEL.search(text)
    if m_base and m_ret:
        b = _parse_money(m_base.group(1))
        r = _parse_money(m_ret.group(1))
        if b > r > 0 and 0.001 < r / b < 0.35:
            return b, r

    # 15. Fallback numérico — construye lista negra de NITs encontrados en texto
    #     para no confundir NIT del emisor/receptor con base sujeta a retención.
    nit_blacklist: set[float] = set()
    # Añadir el NIT conocido del emisor (pasado como hint) para evitar confusiones
    if nit_hint:
        nit_clean = re.sub(r"\D", "", nit_hint)
        if nit_clean:
            nit_blacklist.add(float(nit_clean))
    for pat in (
        # NIT del retenedor (emisor del certificado)
        r"N\.?I\.?T\.?\s*[.:\s#oO]*([0-9]{6,12}|[0-9]{1,3}(?:[.,][0-9]{3})+)",
        r"NIT\s+o\s+C[eé]dula\s*[.:\s]*([0-9]{6,12}|[0-9]{1,3}(?:[.,][0-9]{3})+)",
        r"C\.?C\.?\s*[.:\s#]*([0-9]{5,12})",
        r"[Rr]etenedor\s*[.:\s#]*([0-9]{6,12})",
        r"[Ii]dentificaci[oó]n\s*[.:\s#]*([0-9]{6,12})",
        r"Vendor.{0,15}?Tax\s+No\.?\s*\n?([0-9]{7,12})",
        r"Nit\s+o\s+C\.C\.\s*\n([0-9]{6,12})",
        # NIT del retenido (receptor — NO es la base sujeta a retención)
        r"[Ii]dentificado(?:\s+con)?\s+(?:NIT\s*)?([0-9]{6,12})",
        r"IDENTIFICACI[OÓ]N\s*[/\-]?\s*NIT\s*[:\s]*(?:CC\s+)?([0-9]{6,12})",
        r"(?:C\.?C\.?|NIT)\s+o\s+NIT\s+([0-9]{6,12})",
        r"[Pp]racticada?\s+[Aa]\s*:?[^\n]{0,80}\n[^\n]{0,30}([0-9]{9,10})\b",
        r"[Rr]etenido\s+[Aa][:\s]+[^\n]{0,80}\n[^\n]{0,20}([0-9]{9,10})\b",
        r"\bNr\.?\s*Identificaci[oó]n\s*:\s*([0-9]{7,12})",
        # Números en formato NIT colombiano: 9-12 dígitos seguidos de guión y dígito verificador
        r"\b([0-9]{7,12})-[0-9]\b",
        # Teléfonos y fax — no son montos
        r"(?:Tel[eé]fonos?|Telf?|Celular?|Fax|PBX|Móvil|Movil)[:\s.]*([0-9]{7,12})",
        r"Tel\s+([0-9]{7,12})",
        r"\+57\s*([0-9]{10})",
    ):
        for mt in re.finditer(pat, text, re.I):
            nit_blacklist.add(_parse_money(mt.group(1)))

    nums = [_parse_money(n) for n in re.findall(r"[\d.,]{5,}", text)]
    candidates = [n for n in nums if 1_000 < n < 50_000_000_000 and n not in nit_blacklist]
    for i in range(len(candidates) - 1):
        b, r = candidates[i], candidates[i + 1]
        if b > r > 0 and 0.003 < r / b < 0.25:
            return b, r

    return 0.0, 0.0


# ── Ciudad de retención ───────────────────────────────────────────────────────

_CIUDAD_OCR_FIXES = {
    "MEDELLN": "MEDELLÍN", "MEDELL N": "MEDELLÍN", "BOGOT ": "BOGOTÁ",
    "BOGOT": "BOGOTÁ", "BARRANQUILA": "BARRANQUILLA",
}

_TRAILING_NOISE = frozenset({
    "de", "la", "el", "los", "las", "del", "en", "al", "a", "y", "o", "un", "una",
    "por", "con", "para", "que", "se",
})

def _clean_city(raw: str) -> str:
    """Limpia ciudad: quita prefijos de ruido, valida contra blacklist."""
    city = raw.split("\n")[0].strip()
    city = re.split(r"[,;/]", city)[0].strip()
    city = re.sub(r"^(?:DE\s*:?\s*|:\s*)", "", city, flags=re.I).strip()
    city = re.sub(
        r"(?i)\s*\b(se expide|donde se|esta cert|ificado|art\.|decreto|d\.c\.|s\.a\.s|\|)\b.*$",
        "", city,
    ).strip()
    city = re.sub(r"[\._\|\-]{2,}.*$", "", city).strip()
    city = re.sub(r"\s*[–—]\s*.*$", "", city).strip()  # em/en dash y lo que sigue
    # Strip trailing articles/prepositions and blacklisted words
    words = city.split()
    while words and (
        words[-1].lower().rstrip(".,:;") in _CIUDAD_BLACKLIST
        or words[-1].lower().rstrip(".,:;") in _TRAILING_NOISE
    ):
        words.pop()
    city = " ".join(words)
    # Máximo 3 palabras
    words = city.split()
    city = " ".join(words[:3]) if len(words) > 3 else city
    # Correcciones OCR comunes
    city_up = city.upper().strip()
    city = _CIUDAD_OCR_FIXES.get(city_up, city)
    # Debe comenzar con letra mayúscula
    if city and city[0].islower():
        return ""
    # Descartar si alguna palabra significativa (>3 chars) es del blacklist
    content_words = [w.lower().rstrip(".,:;") for w in city.split() if len(w) > 3]
    if content_words and any(w in _CIUDAD_BLACKLIST for w in content_words):
        return ""
    return city if len(city) >= 3 else ""


def _extract_ciudad_retencion(text: str) -> str:
    # 1. Patrones explícitos de ciudad de retención/consignación
    for pat in (_RE_CIUDAD_PRACT, _RE_CIUDAD_CONSIG):
        m = pat.search(text)
        if m:
            city = _clean_city(m.group(1))
            if len(city) >= 3:
                return city

    # 2. "practicada/o en Medellín"
    m = _RE_CIUDAD_PRACTICADA.search(text)
    if m:
        city = _clean_city(m.group(1))
        if len(city) >= 3:
            return city

    # 3. "en la ciudad de GIRARDOTA"
    m = _RE_CIUDAD_EN.search(text)
    if m:
        city = _clean_city(m.group(1))
        if len(city) >= 3:
            return city

    # 4. "Lugar de expedición: CIUDAD"
    m = _RE_CIUDAD_LUGAR.search(text)
    if m:
        city = _clean_city(m.group(1))
        if len(city) >= 3:
            return city

    # 5. "municipio de Sabaneta"
    m = _RE_CIUDAD_MUNICIPIO.search(text)
    if m:
        city = _clean_city(m.group(1))
        if len(city) >= 3:
            return city

    # 6. Línea de fecha: "Medellín, 12 de enero de 2025" o "MEDELLÍN, 12 ENERO 2025"
    for pat in (_RE_CIUDAD_FECHA, _RE_CIUDAD_FECHA2):
        m = pat.search(text)
        if m:
            city = _clean_city(m.group(1))
            if len(city) >= 3:
                return city

    # 7. "Ciudad: Sincelejo" en encabezado del emisor
    m = _RE_CIUDAD_EMISOR.search(text)
    if m:
        city = _clean_city(m.group(1))
        if len(city) >= 3:
            return city

    return ""


def _fix_pct_as_amount(b: float, r: float) -> float:
    """Si r parece un porcentaje en lugar de valor monetario (ej. 2.5 en vez de 28.928), recalcula."""
    if b > 10_000 and 0 < r <= 30 and r / b < 0.001:
        return round(b * r / 100, 2)
    return r


# ── Entry point ───────────────────────────────────────────────────────────────

def _base_row(path: Path, error: str = "") -> dict:
    """Dict base con todos los campos requeridos."""
    return {
        "razon_social": "", "nit": "", "dv": "", "tipo_doc": "31",
        # Persona natural (tipo_doc=13)
        "primer_apellido": "", "segundo_apellido": "",
        "primer_nombre": "", "otros_nombres": "",
        "direccion": "", "ciudad_retencion": "", "concepto": "", "base": 0.0,
        "retencion": 0.0, "porcentaje": 0.0, "tipo_cert": "",
        "validacion_tasa": 0.0,
        "es_escaneado": False,
        "fuente": "PDF", "archivo": path.name, "error": error,
    }


_MAX_PAGES = 10  # Certificados de retención no superan 10 páginas


def _read_pdf(path: Path) -> tuple[str, bool]:
    """Lee PDF. Si no hay texto extraíble, intenta OCR vía pdfplumber page.to_image()."""
    parts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages[:_MAX_PAGES]:
            t = page.extract_text()
            if t:
                parts.append(t)
    text = "\n".join(parts)
    # Requiere al menos 20 caracteres alfanuméricos — evita tratar PDFs escaneados
    # que solo producen un carácter de artefacto (e.g. ")") como texto válido.
    if len(re.sub(r"\W+", "", text)) >= 20:
        return text, False
    # PDF sin texto (o casi sin texto) → intentar OCR renderizando cada página como imagen
    return _ocr_pdf_pages(path)


def _ocr_pdf_pages(path: Path) -> tuple[str, bool]:
    """OCR sobre páginas de PDF escaneado usando pdfplumber + pytesseract."""
    try:
        import pytesseract
        parts: list[str] = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages[:_MAX_PAGES]:
                img = page.to_image(resolution=200).original
                t = pytesseract.image_to_string(img, lang="spa")
                if t.strip():
                    parts.append(t)
        text = "\n".join(parts)
        return text, not text.strip()
    except Exception:
        return "", True  # Tesseract no disponible o PDF corrupto


def _read_docx(path: Path) -> tuple[str, bool]:
    from docx import Document
    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts), False


def _read_excel(path: Path) -> tuple[str, bool]:
    """Convierte celdas de Excel a texto plano. Soporta .xlsx (openpyxl) y .xls (xlrd)."""
    if path.suffix.lower() == ".xls":
        try:
            import xlrd
            wb = xlrd.open_workbook(str(path))
            parts: list[str] = []
            for ws in wb.sheets():
                for row_idx in range(ws.nrows):
                    def _xls_cell(v: object) -> str:
                        if isinstance(v, float) and v == int(v):
                            return str(int(v))
                        return str(v) if v != "" else ""
                    row_txt = "\t".join(_xls_cell(ws.cell_value(row_idx, col)) for col in range(ws.ncols))
                    if row_txt.strip():
                        parts.append(row_txt)
            return "\n".join(parts), False
        except ImportError:
            raise RuntimeError("Formato .xls no soportado en este entorno. Convierte el archivo a .xlsx e inténtalo de nuevo.")
    import openpyxl
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    parts: list[str] = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            row_txt = "\t".join(str(c) if c is not None else "" for c in row)
            if row_txt.strip():
                parts.append(row_txt)
    wb.close()
    return "\n".join(parts), False


def _preprocess_for_ocr(img: "Image.Image") -> list["Image.Image"]:
    """
    Genera variantes de preprocesamiento para mejorar el OCR en imágenes
    de baja calidad (texto muy claro, baja iluminación, foto de documento).
    """
    from PIL import Image, ImageEnhance, ImageOps
    variants: list[Image.Image] = []

    # Asegurar escala mínima 1600px en el eje largo para buena resolución
    w, h = img.size
    long_side = max(w, h)
    scale = max(2, 1600 // long_side) if long_side < 1600 else 1
    if scale > 1:
        img = img.resize((w * scale, h * scale), Image.LANCZOS)

    gray = img.convert("L")

    # 1. Autocontraste + contraste 3× (fondo muy claro / texto muy débil)
    ac = ImageOps.autocontrast(gray, cutoff=2)
    ac3 = ImageEnhance.Contrast(ac).enhance(3.0)
    variants.append(ac3)

    # 2. Binarización tras autocontraste (umbral 160 después del estiramiento)
    bw = ac3.point(lambda x: 0 if x < 160 else 255, "L")
    variants.append(bw)

    # 3. Original escalado (fallback)
    variants.append(gray)

    return variants


def _read_image(path: Path) -> tuple[str, bool]:
    """
    OCR sobre imagen (JPG/PNG/TIFF) con pytesseract.
    Intenta múltiples preprocesados y devuelve el que extrae más texto.
    """
    try:
        from PIL import Image
        import pytesseract
    except ImportError:
        return "", True  # sin pytesseract → tratar como escaneado

    img = Image.open(path).convert("RGB")
    variants = _preprocess_for_ocr(img)
    img.close()

    best = ""
    configs = ["--oem 3 --psm 6", "--oem 3 --psm 11"]
    for variant in variants:
        for cfg in configs:
            try:
                t = pytesseract.image_to_string(variant, lang="spa", config=cfg)
            except Exception:
                continue
            if len(re.sub(r"\W+", "", t)) > len(re.sub(r"\W+", "", best)):
                best = t
    return best, not best.strip()


def _read_text(path: Path) -> tuple[str, bool]:
    """
    Dispatcher multi-formato. Retorna (texto, es_escaneado).
    Soporta: PDF, DOCX/DOC, XLSX/XLS, JPG/PNG/TIFF/BMP/WEBP.
    """
    suf = path.suffix.lower()
    if suf in (".docx", ".doc"):
        return _read_docx(path)
    if suf in (".xlsx", ".xls"):
        return _read_excel(path)
    if suf in (".jpg", ".jpeg", ".png", ".tiff", ".tif", ".bmp", ".webp"):
        return _read_image(path)
    # Default: PDF
    return _read_pdf(path)


def extract_many(path: "str | Path") -> list[dict]:
    """
    Extrae TODOS los registros de un certificado de retención.
    Un certificado puede tener múltiples conceptos (1302, 1303) o
    múltiples terceros, generando varias filas.

    Retorna lista de dicts con campos:
        razon_social, nit, dv, tipo_doc, direccion, ciudad_retencion,
        concepto, base, retencion, porcentaje, validacion_tasa,
        es_escaneado, tipo_cert, fuente, archivo, error
    """
    path = Path(path)
    base = _base_row(path)

    try:
        text, es_escaneado = _read_text(path)
    except Exception as e:
        base["error"] = f"No se pudo abrir: {e}"
        return [base]

    if es_escaneado:
        base["error"] = "PDF escaneado (imagen) — texto no extraíble"
        base["es_escaneado"] = True
        return [base]

    tipo_cert      = _cert_type(text)
    razon, nit, dv = _extract_emisor(text)
    ciudad         = _extract_ciudad_retencion(text)
    direccion      = _extract_direccion(text)
    tipo_doc       = _tipo_doc_from_context(nit, text)

    # Para persona natural (tipo_doc=13) extraer apellidos y nombres
    ap1 = ap2 = nom1 = otros = ""
    if tipo_doc == "13" and razon:
        ap1, ap2, nom1, otros = _split_persona_natural(razon)

    suf = path.suffix.lower()
    _fuente_map = {
        ".pdf": "PDF", ".docx": "DOCX", ".doc": "DOC",
        ".xlsx": "EXCEL", ".xls": "EXCEL",
        ".jpg": "IMAGEN", ".jpeg": "IMAGEN", ".png": "IMAGEN",
        ".tiff": "IMAGEN", ".tif": "IMAGEN", ".bmp": "IMAGEN", ".webp": "IMAGEN",
    }
    fuente = _fuente_map.get(suf, "PDF")

    def _make_row(concepto: str, porcentaje: float, b: float, r: float) -> dict:
        tasa_calc = round(r / b * 100, 2) if b else 0.0
        return {
            "razon_social":      razon if tipo_doc == "31" else "",
            "nit":               nit,
            "dv":                dv,
            "tipo_doc":          tipo_doc,
            "primer_apellido":   ap1,
            "segundo_apellido":  ap2,
            "primer_nombre":     nom1,
            "otros_nombres":     otros,
            "direccion":         direccion,
            "ciudad_retencion":  ciudad,
            "concepto":          concepto,
            "base":              b,
            "retencion":         r,
            "porcentaje":        porcentaje,
            "validacion_tasa":   tasa_calc,
            "es_escaneado":      False,
            "tipo_cert":         tipo_cert,
            "fuente":            fuente,
            "archivo":           path.name,
            "error":             "",
        }

    # ICA → fila única (va a tabla separada)
    if tipo_cert == "ICA":
        b, r = _extract_amounts(text, tipo_cert, nit_hint=nit)
        r = _fix_pct_as_amount(b, r)
        return [_make_row("ICA", 0.0, b, r)]

    # Intentar múltiples filas de concepto
    multi = _extract_concepto_rows(text, tipo_cert)
    if multi:
        return [_make_row(concepto, pct, b, _fix_pct_as_amount(b, r)) for concepto, pct, b, r in multi]

    # Fallback: fila única
    concepto, porcentaje = _detect_concepto(text, tipo_cert)
    b, r                 = _extract_amounts(text, tipo_cert, nit_hint=nit)
    r = _fix_pct_as_amount(b, r)
    if b == 0 and r == 0:
        base["error"] = "No se encontraron montos"
        base.update({"razon_social": razon if tipo_doc == "31" else "",
                     "nit": nit, "dv": dv, "tipo_doc": tipo_doc,
                     "primer_apellido": ap1, "segundo_apellido": ap2,
                     "primer_nombre": nom1, "otros_nombres": otros,
                     "ciudad_retencion": ciudad, "tipo_cert": tipo_cert})
        return [base]
    return [_make_row(concepto, porcentaje, b, r)]


def extract_one(path: "str | Path") -> dict:
    """
    Extrae el primer registro de un certificado de retención PDF.
    Alias de extract_many()[0] — mantiene compatibilidad hacia atrás.

    Para certificados con múltiples conceptos usar extract_many().
    """
    rows = extract_many(path)
    return rows[0] if rows else _base_row(Path(path), error="Sin resultados")
